"""
Document Processing Service
Handles PDF, TXT, Markdown, CSV, DOCX ingestion, text extraction, and chunking
"""

import hashlib
import os
import re
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple

import pdfminer.high_level
from pdfminer.layout import LAParams

# For DOCX
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# For Markdown
try:
    import mistletoe
    HAS_MISTLETOE = True
except ImportError:
    HAS_MISTLETOE = False

from .embedding_service import EmbeddingService
from ..models.database import Document, DocumentChunk, Embedding


class DocumentProcessor:
    """Processes various document types and prepares chunks for embedding."""

    def __init__(self, chunk_size: int = 1024, chunk_overlap: int = 128):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_service = EmbeddingService()

    def process_document(
        self,
        file_path: str,
        file_type: Optional[str] = None,
        project_id: Optional[str] = None
    ) -> Tuple[Document, List[DocumentChunk]]:
        """Process a document and return the document with its chunks."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect file type if not provided
        if file_type is None:
            file_type = self._detect_file_type(path)

        # Extract text based on file type
        text_content = self._extract_text(str(path), file_type)
        checksum = self._calculate_checksum(str(path))
        file_size = path.stat().st_size
        word_count = len(text_content.split())
        page_count = self._estimate_pages(text_content, file_type)

        # Generate chunks
        chunks = self._chunk_text(text_content)

        # Create document record (would be saved to DB)
        import uuid
        from datetime import datetime, timezone
        doc = Document(
            id=str(uuid.uuid4()),
            project_id=project_id,
            filename=path.name,
            file_path=str(path.absolute()),
            file_type=file_type,
            file_size=file_size,
            title=path.stem,
            created_at=datetime.now(timezone.utc),
            processed_at=datetime.now(timezone.utc),
            status="completed",
            page_count=page_count,
            word_count=word_count,
            checksum=checksum,
        )

        # Create chunks (would be saved to DB)
        document_chunks = []
        for i, chunk_text in enumerate(chunks):
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                document_id=doc.id,
                content=chunk_text,
                chunk_index=i,
                start_char=sum(len(chunks[j]) for j in range(i)),
                metadata={"chunk_size": len(chunk_text)},
            )
            document_chunks.append(chunk)

        return doc, document_chunks

    def _detect_file_type(self, path: Path) -> str:
        """Detect file type from extension."""
        ext = path.suffix.lower()
        mapping = {
            '.pdf': 'pdf',
            '.txt': 'txt',
            '.md': 'md',
            '.csv': 'csv',
            '.docx': 'docx',
            '.doc': 'docx',
        }
        return mapping.get(ext, 'txt')

    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from various file formats."""
        if file_type == 'pdf':
            return self._extract_pdf(file_path)
        elif file_type == 'docx' and HAS_DOCX:
            return self._extract_docx(file_path)
        elif file_type == 'csv':
            return self._extract_csv(file_path)
        else:
            # TXT, MD, and fallback
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfminer."""
        laparams = LAParams()
        text = pdfminer.high_level.extract_text(
            file_path,
            laparams=laparams,
            maxpages=0  # Extract all pages
        )
        return text

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return '\n'.join(paragraphs)

    def _extract_csv(self, file_path: str) -> str:
        """Extract text representation from CSV."""
        import csv
        rows = []
        with open(file_path, 'r', newline='', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(', '.join(str(cell) for cell in row))
        return '\n'.join(rows)

    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into overlapping segments."""
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        words = text.split()
        current_chunk_words = []
        current_length = 0

        for word in words:
            word_length = len(word)
            if current_length + word_length + 1 > self.chunk_size:
                # Store current chunk
                chunks.append(' '.join(current_chunk_words))
                # Start new chunk with overlap
                overlap_words = self.chunk_overlap // 10  # Approximate word overlap
                current_chunk_words = current_chunk_words[-overlap_words:] if overlap_words < len(current_chunk_words) else []
                current_length = sum(len(w) for w in current_chunk_words) + max(0, len(current_chunk_words) - 1)

            current_chunk_words.append(word)
            current_length += word_length

        # Add final chunk
        if current_chunk_words:
            chunks.append(' '.join(current_chunk_words))

        return chunks

    def _calculate_checksum(self, file_path: str) -> str:
        """Calculate SHA-256 checksum of file."""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _estimate_pages(self, text: str, file_type: str) -> Optional[int]:
        """Estimate page count from text or file type."""
        if file_type == 'pdf' and '\x0c' in open('/dev/null').read():  # Can't easily detect without parsing
            pass  # Would use PDF parser metadata
        # Rough estimate: ~3000 chars per page
        return max(1, len(text) // 3000)


class EmbeddingService:
    """Service for generating text embeddings."""

    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        self.model_name = model_name
        try:
            from sentence_transformers import SentenceTransformer
            self.model = SentenceTransformer(model_name)
            self.has_model = True
        except ImportError:
            self.model = None
            self.has_model = False

    def generate_embedding(self, text: str) -> Optional[List[float]]:
        """Generate embedding for a text chunk."""
        if not self.has_model:
            return None
        embedding = self.model.encode(text, show_progress_bar=False)
        return embedding.tolist()